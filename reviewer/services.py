from __future__ import annotations
import asyncio
import io
import re
import time
import logging
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from typing import Iterable, List, Sequence, Tuple, Optional, Callable, Any
from django.conf import settings
from django.db import transaction
from docx import Document
from openai import OpenAI, AsyncOpenAI
from serpapi import GoogleSearch
from asgiref.sync import sync_to_async
from .models import DocumentChunk, FileUploadBatch, UploadedFile


logger = logging.getLogger(__name__)

DEFAULT_EMBEDDING_MODEL = "text-embedding-3-small"
DEFAULT_COMPLETION_MODEL = "gpt-4.1"
MAX_CHUNK_CHAR_LENGTH = 4000
# OpenAI embedding limits: 8191 tokens per text, batch up to ~2000 texts
# For large documents, process embeddings in smaller batches to avoid rate limits
EMBEDDING_BATCH_SIZE = 100  # Number of texts to embed at once (increased for speed)
EMBEDDING_RETRY_DELAY = 1  # Seconds to wait between retries (reduced)
MAX_EMBEDDING_RETRIES = 5  # Maximum retry attempts for rate limits
MAX_CONCURRENT_EMBEDDINGS = 5  # Maximum concurrent embedding API calls


@dataclass
class RetrievedChunk:
    chunk: DocumentChunk
    similarity: float


def _get_openai_client() -> OpenAI:
    api_key = settings.OPENAI_API_KEY
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY env var must be set before calling OpenAI APIs.")
    return OpenAI(api_key=api_key)


def _get_async_openai_client() -> AsyncOpenAI:
    """Get an async OpenAI client for concurrent operations."""
    api_key = settings.OPENAI_API_KEY
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY env var must be set before calling OpenAI APIs.")
    return AsyncOpenAI(api_key=api_key)


def _extract_text_segments(file_obj) -> List[str]:
    """
    Read a DOCX file-like object and return non-empty text segments.
    """
    document = Document(file_obj)
    segments = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            segments.append(text)
    return segments


def _batch_segments(segments: Sequence[str], max_chars: int = MAX_CHUNK_CHAR_LENGTH) -> List[str]:
    """
    Merge sequential segments into chunks that respect the max_chars limit.
    """
    batches: List[str] = []
    current: List[str] = []
    current_len = 0

    for segment in segments:
        if not segment:
            continue
        candidate_len = current_len + len(segment) + (1 if current else 0)
        if candidate_len > max_chars and current:
            batches.append("\n".join(current))
            current = [segment]
            current_len = len(segment)
        else:
            current.append(segment)
            current_len = candidate_len

    if current:
        batches.append("\n".join(current))

    return batches


def _embed_texts(texts: Sequence[str], *, model: str = DEFAULT_EMBEDDING_MODEL) -> List[List[float]]:
    client = _get_openai_client()
    response = client.embeddings.create(model=model, input=list(texts))
    return [item.embedding for item in response.data]


def _embed_texts_chunked(
    texts: Sequence[str],
    *,
    model: str = DEFAULT_EMBEDDING_MODEL,
    batch_size: int = EMBEDDING_BATCH_SIZE,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> List[List[float]]:
    """
    Embed texts in smaller batches to handle large documents and avoid rate limits.

    Args:
        texts: List of text segments to embed
        model: OpenAI embedding model to use
        batch_size: Number of texts to embed per API call
        progress_callback: Optional callback(processed, total) for progress tracking

    Returns:
        List of embedding vectors in the same order as input texts
    """
    if not texts:
        return []

    client = _get_openai_client()
    all_embeddings: List[List[float]] = []
    total_texts = len(texts)

    for i in range(0, total_texts, batch_size):
        batch = texts[i:i + batch_size]

        # Retry logic for rate limits
        for attempt in range(MAX_EMBEDDING_RETRIES):
            try:
                response = client.embeddings.create(model=model, input=list(batch))
                batch_embeddings = [item.embedding for item in response.data]
                all_embeddings.extend(batch_embeddings)

                if progress_callback:
                    progress_callback(min(i + batch_size, total_texts), total_texts)

                break  # Success, exit retry loop

            except Exception as e:
                error_str = str(e).lower()
                if "rate" in error_str or "limit" in error_str or "429" in error_str:
                    if attempt < MAX_EMBEDDING_RETRIES - 1:
                        wait_time = EMBEDDING_RETRY_DELAY * (attempt + 1)
                        logger.warning(f"Rate limit hit, waiting {wait_time}s before retry...")
                        time.sleep(wait_time)
                    else:
                        raise RuntimeError(f"Rate limit exceeded after {MAX_EMBEDDING_RETRIES} retries: {e}")
                else:
                    raise RuntimeError(f"Embedding API error: {e}")

        # Small delay between batches to avoid rate limits
        if i + batch_size < total_texts:
            time.sleep(0.1)

    return all_embeddings


def ingest_docx(
    *,
    file_obj,
    document_type: DocumentChunk.DocumentType,
    title: str | None = None,
    replace_existing: bool = False,
    use_chunked_embedding: bool = True,
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> List[DocumentChunk]:
    """
    Parse, chunk, embed, and persist DOCX content for the provided document type.

    Args:
        file_obj: File-like object containing the DOCX document
        document_type: Type of document (guideline or example)
        title: Optional title override
        replace_existing: Whether to replace existing documents with same title
        use_chunked_embedding: Use chunked embedding for large documents (default True)
        progress_callback: Optional callback(processed, total) for progress tracking

    Returns:
        List of created DocumentChunk objects
    """
    source_name = getattr(file_obj, "name", "") or ""
    file_obj.seek(0)
    segments = _extract_text_segments(file_obj)
    if not segments:
        raise ValueError("لم يتم العثور على نص داخل الملف المرفوع.")

    # Determine the final title
    final_title = title or source_name or document_type

    # Check if a document with this title already exists
    existing_chunks = DocumentChunk.objects.filter(
        document_type=document_type,
        title=final_title
    ).exists()

    if existing_chunks and not replace_existing:
        raise ValueError(f"يوجد مستند بنفس العنوان '{final_title}' بالفعل. الرجاء اختيار عنوان آخر أو حذف المستند الموجود أولاً.")

    batches = _batch_segments(segments)

    # Use chunked embedding for large documents (> 10 batches)
    if use_chunked_embedding and len(batches) > 10:
        embeddings = _embed_texts_chunked(batches, progress_callback=progress_callback)
    else:
        embeddings = _embed_texts(batches)

    if len(batches) != len(embeddings):
        raise RuntimeError("Embedding count mismatch while processing DOCX.")

    with transaction.atomic():
        if replace_existing:
            DocumentChunk.objects.filter(document_type=document_type, title=final_title).delete()

        created_chunks: List[DocumentChunk] = []
        for idx, (text, embedding) in enumerate(zip(batches, embeddings)):
            chunk = DocumentChunk.objects.create(
                document_type=document_type,
                title=final_title,
                source_name=source_name,
                order=idx,
                text=text,
                embedding=embedding,
                metadata={"segment_count": len(text.splitlines())},
            )
            created_chunks.append(chunk)

    return created_chunks


def ingest_multiple_docx(
    *,
    files: List[Tuple],  # List of (file_obj, optional_title) tuples
    document_type: DocumentChunk.DocumentType,
    replace_existing: bool = False,
) -> FileUploadBatch:
    """
    Process multiple DOCX files in a single batch operation.

    This function handles large documents efficiently by:
    1. Processing files sequentially to manage memory
    2. Using chunked embeddings to handle OpenAI rate limits
    3. Tracking progress at the file and batch level

    Args:
        files: List of (file_object, optional_title) tuples
        document_type: Type of documents (guideline or example)
        replace_existing: Whether to replace existing documents with same titles

    Returns:
        FileUploadBatch object with processing results
    """
    # Create batch record
    batch = FileUploadBatch.objects.create(
        document_type=document_type,
        status=FileUploadBatch.Status.PROCESSING,
        total_files=len(files),
    )

    # Create file records
    file_records: List[UploadedFile] = []
    for file_obj, title in files:
        filename = getattr(file_obj, "name", "") or f"file_{len(file_records)}.docx"
        file_size = 0
        try:
            file_obj.seek(0, 2)  # Seek to end
            file_size = file_obj.tell()
            file_obj.seek(0)  # Reset to beginning
        except Exception:
            pass

        file_record = UploadedFile.objects.create(
            batch=batch,
            filename=filename,
            title=title or "",
            file_size=file_size,
            status=UploadedFile.Status.PENDING,
        )
        file_records.append(file_record)

    total_chunks = 0
    processed = 0
    errors = []

    # Process each file
    for idx, ((file_obj, title), file_record) in enumerate(zip(files, file_records)):
        try:
            file_record.status = UploadedFile.Status.PROCESSING
            file_record.save()

            # Use title from parameter or from file record
            use_title = title or file_record.title or None

            chunks = ingest_docx(
                file_obj=file_obj,
                document_type=document_type,
                title=use_title,
                replace_existing=replace_existing,
                use_chunked_embedding=True,
            )

            file_record.status = UploadedFile.Status.COMPLETED
            file_record.chunks_created = len(chunks)
            file_record.save()

            total_chunks += len(chunks)
            processed += 1

        except Exception as e:
            error_msg = str(e)
            file_record.status = UploadedFile.Status.FAILED
            file_record.error_message = error_msg
            file_record.save()
            errors.append(f"{file_record.filename}: {error_msg}")
            logger.error(f"Error processing {file_record.filename}: {e}")

        # Update batch progress
        batch.processed_files = processed
        batch.total_chunks_created = total_chunks
        batch.save()

    # Update final batch status
    if errors:
        if processed == 0:
            batch.status = FileUploadBatch.Status.FAILED
        else:
            batch.status = FileUploadBatch.Status.COMPLETED  # Partial success
        batch.error_message = "\n".join(errors)
    else:
        batch.status = FileUploadBatch.Status.COMPLETED

    batch.save()
    return batch


def get_batch_status(batch_id: str) -> dict:
    """
    Get the current status of a batch upload.

    Args:
        batch_id: UUID of the batch

    Returns:
        Dictionary with batch status and file details
    """
    try:
        batch = FileUploadBatch.objects.get(id=batch_id)
    except FileUploadBatch.DoesNotExist:
        raise ValueError(f"Batch {batch_id} not found")

    files_info = []
    for f in batch.files.all():
        files_info.append({
            "filename": f.filename,
            "title": f.title,
            "status": f.status,
            "chunks_created": f.chunks_created,
            "file_size": f.file_size,
            "error_message": f.error_message,
        })

    progress = 0.0
    if batch.total_files > 0:
        progress = (batch.processed_files / batch.total_files) * 100

    return {
        "batch_id": str(batch.id),
        "document_type": batch.document_type,
        "status": batch.status,
        "total_files": batch.total_files,
        "processed_files": batch.processed_files,
        "total_chunks_created": batch.total_chunks_created,
        "progress_percentage": round(progress, 2),
        "error_message": batch.error_message,
        "files": files_info,
        "created_at": batch.created_at.isoformat(),
        "updated_at": batch.updated_at.isoformat(),
    }


# =============================================================================
# ASYNC PROCESSING FUNCTIONS FOR BIG DATA HANDLING
# =============================================================================

async def _async_embed_texts(
    texts: Sequence[str],
    *,
    model: str = DEFAULT_EMBEDDING_MODEL,
) -> List[List[float]]:
    """
    Asynchronously embed texts using OpenAI's async client.

    Args:
        texts: List of text segments to embed
        model: OpenAI embedding model to use

    Returns:
        List of embedding vectors in the same order as input texts
    """
    if not texts:
        return []

    client = _get_async_openai_client()
    response = await client.embeddings.create(model=model, input=list(texts))
    return [item.embedding for item in response.data]


async def _async_embed_texts_chunked(
    texts: Sequence[str],
    *,
    model: str = DEFAULT_EMBEDDING_MODEL,
    batch_size: int = EMBEDDING_BATCH_SIZE,
    max_concurrent: int = MAX_CONCURRENT_EMBEDDINGS,
) -> List[List[float]]:
    """
    Asynchronously embed texts in batches with concurrency control.

    This function handles large documents by:
    1. Breaking texts into smaller batches
    2. Processing batches concurrently (with limits to avoid rate limits)
    3. Implementing retry logic for rate limit errors

    Args:
        texts: List of text segments to embed
        model: OpenAI embedding model to use
        batch_size: Number of texts per batch
        max_concurrent: Maximum concurrent API calls

    Returns:
        List of embedding vectors in the same order as input texts
    """
    if not texts:
        return []

    client = _get_async_openai_client()
    all_embeddings: List[Optional[List[float]]] = [None] * len(texts)
    total_texts = len(texts)

    # Create batches with their indices
    batches = []
    for i in range(0, total_texts, batch_size):
        batch_texts = texts[i:i + batch_size]
        batch_indices = list(range(i, min(i + batch_size, total_texts)))
        batches.append((batch_indices, batch_texts))

    semaphore = asyncio.Semaphore(max_concurrent)

    async def process_batch(batch_indices: List[int], batch_texts: List[str]) -> None:
        async with semaphore:
            for attempt in range(MAX_EMBEDDING_RETRIES):
                try:
                    response = await client.embeddings.create(
                        model=model,
                        input=list(batch_texts)
                    )
                    for idx, item in zip(batch_indices, response.data):
                        all_embeddings[idx] = item.embedding
                    return
                except Exception as e:
                    error_str = str(e).lower()
                    if "rate" in error_str or "limit" in error_str or "429" in error_str:
                        if attempt < MAX_EMBEDDING_RETRIES - 1:
                            wait_time = EMBEDDING_RETRY_DELAY * (attempt + 1)
                            logger.warning(f"Rate limit hit, waiting {wait_time}s before retry...")
                            await asyncio.sleep(wait_time)
                        else:
                            raise RuntimeError(f"Rate limit exceeded after {MAX_EMBEDDING_RETRIES} retries: {e}")
                    else:
                        raise RuntimeError(f"Embedding API error: {e}")

    # Process all batches concurrently
    await asyncio.gather(*[
        process_batch(indices, texts_batch)
        for indices, texts_batch in batches
    ])

    return all_embeddings


async def async_ingest_docx(
    *,
    file_content: bytes,
    filename: str,
    document_type: DocumentChunk.DocumentType,
    title: str | None = None,
    replace_existing: bool = False,
) -> List[DocumentChunk]:
    """
    Asynchronously parse, chunk, embed, and persist DOCX content.

    This is the async version of ingest_docx for handling multiple files concurrently.

    Args:
        file_content: Raw bytes of the DOCX file
        filename: Name of the file
        document_type: Type of document (guideline or example)
        title: Optional title override
        replace_existing: Whether to replace existing documents with same title

    Returns:
        List of created DocumentChunk objects
    """
    # Parse DOCX in a thread pool (python-docx is not async)
    loop = asyncio.get_event_loop()

    def extract_segments():
        file_obj = io.BytesIO(file_content)
        document = Document(file_obj)
        segments = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                segments.append(text)
        return segments

    with ThreadPoolExecutor() as executor:
        segments = await loop.run_in_executor(executor, extract_segments)

    if not segments:
        raise ValueError("لم يتم العثور على نص داخل الملف المرفوع.")

    # Determine the final title
    final_title = title or filename or document_type

    # Check if a document with this title already exists (sync DB operation)
    @sync_to_async
    def check_existing():
        return DocumentChunk.objects.filter(
            document_type=document_type,
            title=final_title
        ).exists()

    existing_chunks = await check_existing()

    if existing_chunks and not replace_existing:
        raise ValueError(
            f"يوجد مستند بنفس العنوان '{final_title}' بالفعل. "
            "الرجاء اختيار عنوان آخر أو حذف المستند الموجود أولاً."
        )

    # Batch segments
    batches = _batch_segments(segments)

    # Embed texts asynchronously
    if len(batches) > 10:
        embeddings = await _async_embed_texts_chunked(batches)
    else:
        embeddings = await _async_embed_texts(batches)

    if len(batches) != len(embeddings):
        raise RuntimeError("Embedding count mismatch while processing DOCX.")

    # Save to database (sync operation wrapped in async)
    @sync_to_async
    def save_chunks():
        with transaction.atomic():
            if replace_existing:
                DocumentChunk.objects.filter(
                    document_type=document_type,
                    title=final_title
                ).delete()

            created_chunks: List[DocumentChunk] = []
            for idx, (text, embedding) in enumerate(zip(batches, embeddings)):
                chunk = DocumentChunk.objects.create(
                    document_type=document_type,
                    title=final_title,
                    source_name=filename,
                    order=idx,
                    text=text,
                    embedding=embedding,
                    metadata={"segment_count": len(text.splitlines())},
                )
                created_chunks.append(chunk)
            return created_chunks

    return await save_chunks()


async def async_ingest_multiple_docx(
    *,
    files_data: List[Tuple[bytes, str, Optional[str]]],  # (content, filename, title)
    document_type: DocumentChunk.DocumentType,
    replace_existing: bool = False,
    max_concurrent_files: int = 3,
) -> FileUploadBatch:
    """
    Asynchronously process multiple DOCX files concurrently.

    This function handles big data (multiple large files) efficiently by:
    1. Processing files concurrently with controlled parallelism
    2. Using async embeddings to avoid blocking
    3. Tracking progress at file and batch levels

    Args:
        files_data: List of (file_content_bytes, filename, optional_title) tuples
        document_type: Type of documents (guideline or example)
        replace_existing: Whether to replace existing documents with same titles
        max_concurrent_files: Maximum files to process concurrently

    Returns:
        FileUploadBatch object with processing results
    """
    # Create batch record
    @sync_to_async
    def create_batch():
        return FileUploadBatch.objects.create(
            document_type=document_type,
            status=FileUploadBatch.Status.PROCESSING,
            total_files=len(files_data),
        )

    batch = await create_batch()

    # Create file records
    @sync_to_async
    def create_file_records():
        records = []
        for content, filename, title in files_data:
            record = UploadedFile.objects.create(
                batch=batch,
                filename=filename,
                title=title or "",
                file_size=len(content),
                status=UploadedFile.Status.PENDING,
            )
            records.append(record)
        return records

    file_records = await create_file_records()

    results = {
        "total_chunks": 0,
        "processed": 0,
        "errors": [],
    }

    semaphore = asyncio.Semaphore(max_concurrent_files)

    async def process_file(
        file_data: Tuple[bytes, str, Optional[str]],
        file_record: UploadedFile
    ):
        async with semaphore:
            content, filename, title = file_data

            @sync_to_async
            def update_status(status, chunks=0, error=""):
                file_record.status = status
                if chunks:
                    file_record.chunks_created = chunks
                if error:
                    file_record.error_message = error
                file_record.save()

            try:
                await update_status(UploadedFile.Status.PROCESSING)

                chunks = await async_ingest_docx(
                    file_content=content,
                    filename=filename,
                    document_type=document_type,
                    title=title,
                    replace_existing=replace_existing,
                )

                await update_status(UploadedFile.Status.COMPLETED, len(chunks))
                results["total_chunks"] += len(chunks)
                results["processed"] += 1

            except Exception as e:
                error_msg = str(e)
                await update_status(UploadedFile.Status.FAILED, error=error_msg)
                results["errors"].append(f"{filename}: {error_msg}")
                logger.error(f"Error processing {filename}: {e}")

            # Update batch progress
            @sync_to_async
            def update_batch_progress():
                batch.processed_files = results["processed"]
                batch.total_chunks_created = results["total_chunks"]
                batch.save()

            await update_batch_progress()

    # Process all files concurrently
    await asyncio.gather(*[
        process_file(file_data, file_record)
        for file_data, file_record in zip(files_data, file_records)
    ])

    # Update final batch status
    @sync_to_async
    def finalize_batch():
        if results["errors"]:
            if results["processed"] == 0:
                batch.status = FileUploadBatch.Status.FAILED
            else:
                batch.status = FileUploadBatch.Status.COMPLETED  # Partial success
            batch.error_message = "\n".join(results["errors"])
        else:
            batch.status = FileUploadBatch.Status.COMPLETED
        batch.save()

    await finalize_batch()

    # Refresh batch from DB
    @sync_to_async
    def refresh_batch():
        batch.refresh_from_db()
        return batch

    return await refresh_batch()


def _cosine_similarity(vec_a: Sequence[float], vec_b: Sequence[float]) -> float:
    dot_product = 0.0
    mag_a = 0.0
    mag_b = 0.0
    for a, b in zip(vec_a, vec_b):
        dot_product += a * b
        mag_a += a * a
        mag_b += b * b
    if mag_a == 0.0 or mag_b == 0.0:
        return 0.0
    return dot_product / ((mag_a ** 0.5) * (mag_b ** 0.5))


def retrieve_similar_chunks(
    *,
    query_text: str,
    document_type: DocumentChunk.DocumentType,
    limit: int | None = 3,
) -> List[RetrievedChunk]:
    chunks = list(DocumentChunk.objects.filter(document_type=document_type))
    if not chunks:
        return []

    query_embedding = _embed_texts([query_text])[0]
    scored = [
        RetrievedChunk(chunk=chunk, similarity=_cosine_similarity(query_embedding, chunk.embedding))
        for chunk in chunks
    ]
    scored.sort(key=lambda item: item.similarity, reverse=True)
    if limit is None or limit >= len(scored):
        return scored
    return scored[:limit]


def _preprocess_honorifics(text: str) -> str:
    """
    Preprocess text to remove excessive honorifics while keeping official titles.

    IMPORTANT: This function preserves quoted text (text within quotation marks) unchanged.

    KEEP (Official Titles):
    - خادم الحرمين الشريفين (official title for Saudi King)
    - صاحب السمو الملكي (official title for Royal Highness)

    REMOVE (Exaggerated Phrases):
    - جلالة الملك المعظم أيده الله → الملك
    - فخامة الرئيس حفظه الله → الرئيس
    - Prayer phrases: حفظه الله، أيده الله، رعاه الله
    - Exaggerated adjectives: المعظم، الجليل
    """
    # Step 1: Extract and preserve quoted text
    # Match various quotation mark styles: "...", «...», "...", '...'
    quote_patterns = [
        r'"([^"]+)"',      # Standard quotes
        r'«([^»]+)»',      # Arabic quotes
        r'"([^"]+)"',      # Curly double quotes
        r"'([^']+)'",      # Single curly quotes
    ]

    # Store quoted texts with placeholders
    quoted_texts = []
    processed_text = text

    for pattern in quote_patterns:
        matches = re.finditer(pattern, processed_text)
        for match in matches:
            placeholder = f"<<<QUOTE_{len(quoted_texts)}>>>"
            quoted_texts.append(match.group(0))  # Store the full match including quotes
            processed_text = processed_text.replace(match.group(0), placeholder, 1)

    # Define comprehensive replacement patterns (order matters - most specific first)
    replacements = [
        # King honorifics - REMOVE exaggerated parts, keep simple title
        # "جلالة الملك المعظم" → "الملك" (remove exaggeration)
        (r'جلالة\s+الملك\s+المعظم', 'الملك'),
        (r'جلالة\s+الملك', 'الملك'),
        (r'حضرة\s+صاحب\s+الجلالة\s+الملك\s+المعظم', 'الملك'),
        (r'حضرة\s+صاحب\s+الجلالة\s+الملك', 'الملك'),
        (r'صاحب\s+الجلالة\s+الملك\s+المعظم', 'الملك'),
        (r'صاحب\s+الجلالة\s+الملك', 'الملك'),
        (r'الملك\s+المعظم', 'الملك'),

        # Sultan honorifics - REMOVE exaggerated parts
        (r'حضرة\s+صاحب\s+الجلالة\s+السلطان\s+المعظم', 'السلطان'),
        (r'حضرة\s+صاحب\s+الجلالة\s+السلطان', 'السلطان'),
        (r'صاحب\s+الجلالة\s+السلطان\s+المعظم', 'السلطان'),
        (r'صاحب\s+الجلالة\s+السلطان', 'السلطان'),
        (r'جلالة\s+السلطان', 'السلطان'),
        (r'السلطان\s+المعظم', 'السلطان'),

        # President honorifics - REMOVE "فخامة" but this is less common in Gulf news
        (r'فخامة\s+الرئيس', 'الرئيس'),

        # Prince titles - KEEP "صاحب السمو الملكي" as it's an official title
        # But remove exaggerated adjectives like "الجليل" or "الكريم" when standalone
        (r'سموه\s+الكريم', 'سموه'),
        (r'سموه\s+الجليل', 'سموه'),

        # Prayer phrases - DELETE completely (these are pure exaggeration, not titles)
        (r'\s+حفظه\s+الله\s+', ' '),
        (r'\s+حفظها\s+الله\s+', ' '),
        (r'\s+رعاه\s+الله\s+', ' '),
        (r'\s+رعاها\s+الله\s+', ' '),
        (r'\s+نصره\s+الله\s+', ' '),
        (r'\s+أيده\s+الله\s+', ' '),
        (r'\s+أطال\s+الله\s+عمره\s+', ' '),
        (r'\s+أدام\s+الله\s+عزه\s+', ' '),
        (r'\s+حفظهما\s+الله\s+', ' '),
        (r'\s+حفظهم\s+الله\s+', ' '),
        (r'\s+حفظه\s+الله\s*$', ' '),  # At end of sentence
        (r'^\s*حفظه\s+الله\s+', ' '),  # At start of sentence

        # Exaggerated adjectives - DELETE
        (r'\bخالص\s+', ''),  # "خالص تهانيه" → "تهانيه"
        (r'\bالجليل\s+', ''),  # When used as standalone adjective
    ]

    # Step 2: Apply all replacements (to non-quoted text)
    for pattern, replacement in replacements:
        processed_text = re.sub(pattern, replacement, processed_text, flags=re.IGNORECASE)

    # Step 3: Clean up multiple spaces and trim
    processed_text = re.sub(r'\s+', ' ', processed_text)
    processed_text = processed_text.strip()

    # Step 4: Restore quoted texts
    for i, quoted_text in enumerate(quoted_texts):
        placeholder = f"<<<QUOTE_{i}>>>"
        processed_text = processed_text.replace(placeholder, quoted_text)

    return processed_text


def build_review_prompt(news_text: str, guidelines: Iterable[RetrievedChunk], examples: Iterable[RetrievedChunk]) -> List[dict]:
    # Preprocess text to handle honorifics before sending to AI
    processed_news_text = _preprocess_honorifics(news_text)
    
    guideline_section_lines = []
    for idx, item in enumerate(guidelines, start=1):
        guideline_section_lines.append(f"{idx}. {item.chunk.text}")
    guideline_section = "\n".join(guideline_section_lines) if guideline_section_lines else "No guidelines available."

    example_section_lines = []
    for idx, item in enumerate(examples, start=1):
        example_section_lines.append(f"### Example {idx}\n{item.chunk.text}")
    example_section = "\n\n".join(example_section_lines) if example_section_lines else "No examples available."

    user_prompt = (
        "🔴🔴🔴 تعليمات صارمة - اقرأها بعناية! 🔴🔴🔴\n\n"
        "⛔ لا تُعِد صياغة الخبر! ⛔\n"
        "⛔ لا تُغيّر كلمات الخبر الأصلي! ⛔\n"
        "⛔ لا تُضِف أي معلومات جديدة! ⛔\n\n"
        "✅ فقط طبّق هذه التعديلات المحددة:\n"
        "1. أضف الجنسية للمسؤول في العنوان (إذا لم تكن موجودة)\n"
        "2. أضف سطر المدينة والوكالة في البداية (إذا لم يكن موجوداً)\n"
        "3. أزل الألقاب المبالغ فيها (جلالة، فخامة، حفظه الله...)\n"
        "4. أضف (انتهى) في النهاية\n"
        "5. صحح الأخطاء الإملائية فقط\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "🔴 قاعدة الجنسية في العناوين 🔴\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "أضف الجنسية بعد المنصب أو الوزارة مباشرة:\n"
        "• 'رئيس الوزراء' → 'رئيس الوزراء العراقي'\n"
        "• 'الملك' → 'الملك السعودي'\n"
        "• 'الخارجية' → 'الخارجية السعودية' أو 'الخارجية المصرية'\n"
        "• 'وزير السياحة' → 'وزير السياحة السعودي'\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "🔴 قاعدة اختصار أسماء الوزارات - مهم جداً! 🔴\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "استخدم الصيغة المختصرة للوزارات (بدون كلمة 'وزارة'):\n"
        "• 'وزارة الخارجية العراقية' → 'الخارجية العراقية'\n"
        "• 'وزارة الخارجية السعودية' → 'الخارجية السعودية'\n"
        "• 'وزارة الخارجية المصرية' → 'الخارجية المصرية'\n"
        "• 'وزارة الدفاع الباكستانية' → 'الدفاع الباكستانية'\n"
        "• 'وزارة الداخلية الإماراتية' → 'الداخلية الإماراتية'\n"
        "• 'وزارة الصحة السعودية' → 'الصحة السعودية'\n\n"
        "⚠️ استثناء: 'وزارة الحج والعمرة' تبقى كما هي بدون اختصار (لأنها وحيدة في العالم)\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "🔴 مصطلحات القضية الفلسطينية - استخدم المصطلح الصحيح! 🔴\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "استبدل المصطلحات التالية بالشكل الصحيح:\n"
        "• 'سكان القدس' → 'المقدسيون'\n"
        "• 'شرقي القدس' → 'القدس المحتلة'\n"
        "• 'عرب إسرائيل' أو 'عرب اسرائيل' → 'فلسطينيو 48'\n"
        "• 'الجدار العازل' → 'جدار الفصل العنصري'\n"
        "• 'الأراضي المتنازع عليها' → 'الأرض الفلسطينية المحتلة'\n"
        "• 'جيش الدفاع' أو 'جيش الدفاع الإسرائيلي' → 'جيش الاحتلال الإسرائيلي'\n"
        "• 'شعب غزة' → 'المواطنون الفلسطينيون في قطاع غزة'\n"
        "• 'الحكومة الإسرائيلية' → 'حكومة الاحتلال الإسرائيلي'\n\n"
        "كيف تعرف الجنسية؟\n"
        "• واع = عراقي | واس = سعودي | وام = إماراتي | وكونا = كويتي | بنا = بحريني\n"
        "• بغداد = عراقي | الرياض = سعودي | القاهرة = مصري | الكويت = كويتي\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "🔴🔴🔴 الهيكل المطلوب - مهم جداً! 🔴🔴🔴\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "⚠️ العنوان يجب أن يكون في سطر منفصل تماماً! ⚠️\n\n"
        "الهيكل:\n"
        "السطر 1: العنوان (مع الجنسية)\n"
        "السطر 2: [سطر فارغ]\n"
        "السطر 3: المدينة (يونا/الوكالة) - نص الخبر\n"
        "السطر 4: [سطر فارغ]\n"
        "السطر 5: (انتهى)\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📌 مثال عملي - انتبه للأسطر الفارغة!\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "❌ خطأ (العنوان ملتصق بالمتن):\n"
        "رئيس الوزراء العراقي: أنشأنا 15 مصحة بغداد (يونا/واع) - أكد...\n\n"
        "✅ صحيح (العنوان في سطر منفصل):\n"
        "رئيس الوزراء العراقي: أنشأنا 15 مصحة قسرية\n\n"
        "بغداد (يونا/واع) - أكد رئيس الوزراء العراقي محمد شياع السوداني أن الحكومة أنشأت 15 مصحة.\n\n"
        "(انتهى)\n\n"
        "⚠️ لاحظ: بين العنوان وبين سطر المدينة يوجد سطر فارغ (\\n\\n)!\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "🔴🔴🔴 قواعد إلزامية على معالج الأخبار - يجب تطبيقها جميعاً! 🔴🔴🔴\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "⛔ القاعدة 1 - العنوان في سطر مستقل (إلزامي):\n"
        "   • يجب أن يكون العنوان في سطر منفصل تماماً عن باقي الخبر\n"
        "   • يتبعه سطر فارغ ثم المتن\n\n"
        "⛔ القاعدة 2 - كلمة (انتهى) في سطر مستقل (إلزامي):\n"
        "   • يجب أن تكون (انتهى) في آخر سطر منفصل\n"
        "   • يسبقها سطر فارغ\n\n"
        "⛔ القاعدة 3 - ذكر اسم الوكالة المصدر (إلزامي):\n"
        "   • يجب كتابة اسم الوكالة المصدر بعد (يونا/)\n"
        "   • ❌ خطأ: طوباس (يونا) –\n"
        "   • ✅ صحيح: طوباس (يونا/الغد) –\n"
        "   • ✅ صحيح: بغداد (يونا/واع) –\n"
        "   • ✅ صحيح: الرياض (يونا/واس) –\n\n"
        "⛔ القاعدة 4 - طول العنوان (إلزامي):\n"
        "   • يجب ألا يتجاوز العنوان 12 كلمة\n"
        "   • إذا كان أطول، اختصره مع الحفاظ على المعنى\n\n"
        "⛔ القاعدة 5 - طول الفقرة (إلزامي):\n"
        "   • يجب ألا يزيد عدد الكلمات في كل فقرة عن 50 كلمة\n"
        "   • قسّم الفقرات الطويلة إلى فقرات أقصر\n\n"
        "⛔ القاعدة 6 - استبدال 'مراسلنا/مراسلتنا' (إلزامي):\n"
        "   • يجب استبدال 'مراسلنا' أو 'مراسلتنا' باسم الوكالة المصدر الأصلية (وليس يونا!)\n"
        "   • انظر إلى مصدر الخبر في (يونا/XXX) واستخدم XXX\n"
        "   • ❌ خطأ: أفاد مراسلنا\n"
        "   • ❌ خطأ: أفاد مراسل يونا (يونا ليست المصدر الأصلي!)\n"
        "   • ✅ صحيح: أفاد مراسل قناة الغد (إذا كان المصدر الغد)\n"
        "   • ✅ صحيح: أفاد مراسل واع (إذا كان المصدر واع)\n"
        "   • ✅ صحيح: أفاد مراسل واس (إذا كان المصدر واس)\n"
        "   • 🔴 مهم: يونا هي الناقلة وليست المصدر الأصلي!\n\n"
        "⛔ القاعدة 7 - التنويع في الأفعال الإعلامية (إلزامي):\n"
        "   • ممنوع تكرار نفس الفعل الإعلامي أكثر من مرة في الخبر\n"
        "   • إذا استخدمت 'أفاد' في الفقرة الأولى، استخدم فعلاً مختلفاً في الفقرة الثانية\n"
        "   • الأفعال المتاحة: قال، ذكر، أوضح، أكد، شدّد، بيّن، أشار، أفاد، أعلن، أفصح، عبّر، أعرب، نوّه، أبان،\n"
        "     كشف، لفت، رجّح، جزم، نفى، أضاف، تابع، ختم، أورد، نقل، حدّد، دعا، طالب، حذّر، استنكر، ندّد،\n"
        "     رحّب، أبدى، اعتبر، رأى، توقّع، أوصى، طرح، أوعز، وجّه، استعرض، صرّح، أخبر، أبلغ، أطلع، أبرز، أظهر\n"
        "   • مثال:\n"
        "     ❌ خطأ: أفاد مراسل... وأضاف أن...\n"
        "     ✅ صحيح: أفاد مراسل... وأوضح أن... / وبيّن أن... / ولفت إلى أن...\n\n"
        "⛔ القاعدة 8 - تماسك الجمل وأدوات الربط (إلزامي):\n"
        "   • يجب الحفاظ على تماسك الجمل وربطها بسلاسة\n"
        "   • استخدم أدوات الربط المناسبة: و، كما، فيما، بينما، في حين، إضافة إلى ذلك، علاوة على ذلك\n\n"
        "⛔ القاعدة 9 - مصطلحات فلسطين (إلزامي):\n"
        "   • في أخبار فلسطين ممنوع استخدام: مقتل، مصرع، موت، قتيل، قتلى\n"
        "   • يجب استخدام: استشهاد، شهيد، شهداء\n"
        "   • ❌ خطأ: مقتل 5 فلسطينيين\n"
        "   • ✅ صحيح: استشهاد 5 فلسطينيين\n"
        "   • ❌ خطأ: جيش الدفاع الإسرائيلي\n"
        "   • ✅ صحيح: جيش الاحتلال الإسرائيلي\n\n"
        "⛔ القاعدة 10 - صيغة التاريخ (إلزامي):\n"
        "   • يجب كتابة التاريخ بالميلادي فقط بهذه الصيغة: اليوم DD الشهر YYYYم\n"
        "   • ✅ صحيح: الأحد 21 ديسمبر 2025م\n"
        "   • ✅ صحيح: الخميس 5 يناير 2025م\n"
        "   • ❌ خطأ: 21 كانون الأول (شهور سريانية ممنوعة)\n"
        "   • ❌ خطأ: 15 جمادى الأولى 1446هـ (شهور هجرية ممنوعة)\n\n"
        "⛔ القاعدة 11 - الصفات التقريرية (إلزامي):\n"
        "   • ممنوع استخدام الصفات التقريرية بدون دليل\n"
        "   • الصفات الممنوعة: ضخم، غير مسبوق، تاريخي، هائل، كبير جداً\n"
        "   • يُسمح بها فقط إذا كانت مدعومة ببيانات أو أرقام محددة\n\n"
        "⛔ القاعدة 12 - الأقواس (إلزامي):\n"
        "   • تجنب الإفراط في استخدام الأقواس\n"
        "   • استخدم الأقواس فقط عند الضرورة القصوى\n\n"
        "⛔ القاعدة 13 - التكرار (إلزامي):\n"
        "   • ممنوع التكرار غير المبرر للجمل أو العبارات\n"
        "   • إذا ذُكرت معلومة مرة، لا تُكررها\n\n"
        "⛔ القاعدة 14 - علامات الترقيم (إلزامي):\n"
        "   • يجب الاهتمام بعلامات الترقيم (الفواصل والنقاط)\n"
        "   • هذا بنفس أهمية تصحيح الأخطاء الإملائية\n"
        "   • تأكد من وجود فاصلة بين الجمل المعطوفة\n"
        "   • تأكد من وجود نقطة في نهاية كل فقرة\n\n"
        "🚨🚨🚨 تحذير: إذا لم تُطبق أي من هذه القواعد، فأنت قد فشلت! 🚨🚨🚨\n\n"
        "### Editorial Guidelines\n"
        f"{guideline_section}\n\n"
        "### Reference News Examples\n"
        f"{example_section}\n\n"
        "### Article Requiring Review\n"
        f"{processed_news_text}\n\n"
        "FIRST: Validate that the text above is a legitimate news article. "
        "If it is random, inappropriate, meaningless, or not a news article, "
        "respond ONLY with: 'ERROR: النص المقدم غير مناسب أو غير صالح للمعالجة. يرجى تقديم خبر صحيح.'\n\n"
        "⚠️⚠️⚠️ CRITICAL RULE: UNDERSTAND THE DIFFERENCE BETWEEN OFFICIAL TITLES AND EXAGGERATION ⚠️⚠️⚠️\n\n"
        "🔴 ABSOLUTE RULE - READ THIS CAREFULLY:\n"
        "There is a HUGE difference between:\n"
        "1. OFFICIAL STATE TITLES (الألقاب الرسمية للدولة) = These are REAL titles, NOT exaggeration → MUST KEEP\n"
        "2. EXAGGERATED PHRASES (التفخيم والتعظيم) = These are praise phrases, NOT titles → MUST REMOVE\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "✅ WHAT TO KEEP - THESE ARE OFFICIAL TITLES (DO NOT TOUCH!):\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "1. ✅ 'خادم الحرمين الشريفين' - Official title of Saudi King (like a last name)\n"
        "   Why? This is his OFFICIAL STATE TITLE, not exaggeration!\n"
        "   ⚠️ NEVER remove this! It's like removing someone's official job title!\n\n"
        "2. ✅ 'صاحب السمو الملكي' - Official Royal Highness title (government-recognized)\n"
        "   Why? This is the OFFICIAL PROTOCOL title for princes, grandson of the King!\n"
        "   ⚠️ NEVER remove this! It's their official designation in the state!\n\n"
        "3. ✅ 'ولي العهد' - Official Crown Prince position\n"
        "4. ✅ 'رئيس مجلس الوزراء' - Official Prime Minister position\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "❌ WHAT TO REMOVE - THESE ARE EXAGGERATION (DELETE OR SIMPLIFY!):\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "1. ❌ 'جلالة الملك المعظم أيده الله' → ✅ 'الملك'\n"
        "   Why? 'جلالة' and 'المعظم' and 'أيده الله' are EXAGGERATION, not official titles!\n\n"
        "2. ❌ 'جلالة الملك' → ✅ 'الملك'\n"
        "3. ❌ 'حضرة صاحب الجلالة الملك' → ✅ 'الملك'\n"
        "4. ❌ 'صاحب الجلالة السلطان' → ✅ 'السلطان'\n"
        "5. ❌ 'فخامة الرئيس' → ✅ 'الرئيس'\n"
        "6. ❌ Prayer phrases: 'حفظه الله', 'أيده الله', 'رعاه الله', 'نصره الله' → DELETE COMPLETELY\n"
        "7. ❌ Exaggeration words: 'المعظم', 'الجليل', 'خالص' → DELETE\n"
        "8. ❌ 'سموه الكريم' → ✅ 'سموه'\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📌 KEY DISTINCTION (READ THIS 10 TIMES!):\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "• 'خادم الحرمين الشريفين' = Official title (like saying 'Dr.' or 'President') → KEEP!\n"
        "• 'صاحب السمو الملكي' = Official royal protocol title → KEEP!\n"
        "• 'جلالة الملك المعظم' = Exaggerated praise → REMOVE, simplify to 'الملك'\n"
        "• 'حفظه الله' 'أيده الله' = Prayer/supplication → DELETE COMPLETELY\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📚 MANDATORY EXAMPLES - STUDY THESE CAREFULLY:\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "Example 1 - Removing exaggeration, keeping simple title:\n"
        "❌ BEFORE: 'بعث جلالة الملك المعظم حمد بن عيسى آل خليفة حفظه الله برقية تهنئة خالصة'\n"
        "✅ AFTER:  'بعث الملك حمد بن عيسى آل خليفة برقية تهنئة'\n"
        "What we removed: جلالة (exaggeration), المعظم (exaggeration), حفظه الله (prayer), خالصة (exaggeration)\n"
        "What we kept: الملك (simple title)\n\n"
        "Example 2 - KEEPING official 'صاحب السمو الملكي' because it's OFFICIAL:\n"
        "❌ BEFORE: 'استقبل صاحب السمو الملكي الأمير سلمان بن حمد آل خليفة ولي العهد رئيس مجلس الوزراء حفظه الله'\n"
        "✅ AFTER:  'استقبل صاحب السمو الملكي الأمير سلمان بن حمد آل خليفة ولي العهد رئيس مجلس الوزراء'\n"
        "What we removed: حفظه الله (prayer phrase only!)\n"
        "What we kept: صاحب السمو الملكي (OFFICIAL TITLE - DO NOT TOUCH!), الأمير, ولي العهد, رئيس مجلس الوزراء\n"
        "⚠️ CRITICAL: We did NOT remove 'صاحب السمو الملكي' because it is an OFFICIAL STATE TITLE!\n\n"
        "Example 3 - KEEPING official 'خادم الحرمين الشريفين' because it's OFFICIAL:\n"
        "❌ BEFORE: 'خادم الحرمين الشريفين الملك سلمان بن عبدالعزيز آل سعود حفظه الله أيده الله'\n"
        "✅ AFTER:  'خادم الحرمين الشريفين الملك سلمان بن عبدالعزيز آل سعود'\n"
        "What we removed: حفظه الله (prayer), أيده الله (prayer)\n"
        "What we kept: خادم الحرمين الشريفين (OFFICIAL SAUDI KING TITLE - NEVER REMOVE!), الملك\n"
        "⚠️ CRITICAL: We did NOT remove 'خادم الحرمين الشريفين' because it is the OFFICIAL TITLE of Saudi King!\n\n"
        "Example 4 - What happens when we see 'جلالة' vs 'صاحب السمو الملكي':\n"
        "❌ WRONG:  'جلالة الملك' → 'جلالة الملك' (keeping exaggeration)\n"
        "✅ RIGHT:  'جلالة الملك' → 'الملك' (removed exaggeration)\n"
        "❌ WRONG:  'صاحب السمو الملكي الأمير' → 'الأمير' (removed official title!)\n"
        "✅ RIGHT:  'صاحب السمو الملكي الأمير' → 'صاحب السمو الملكي الأمير' (kept official title!)\n\n"
        "🔴 FINAL WARNING:\n"
        "If you remove 'خادم الحرمين الشريفين' or 'صاحب السمو الملكي', you have FAILED!\n"
        "These are OFFICIAL STATE TITLES, not exaggeration!\n\n"
        "IMPORTANT: When replacing exaggerated titles, preserve 'ال' (definite article):\n"
        "✅ CORRECT: 'السلطان' (with ال), 'الملك' (with ال)\n"
        "❌ WRONG: 'سلطان' (without ال), 'ملك' (without ال)\n\n"
        "STEP-BY-STEP PROCESS:\n"
        "1. IMPORTANT: DO NOT modify any text inside quotation marks (\"...\", «...», \"...\")!\n"
        "   Quoted text must remain EXACTLY as it appears, including any honorifics or titles.\n"
        "   Example: If someone said \"جلالة الملك المعظم\" in quotes, keep it unchanged!\n"
        "2. Scan the entire article for ALL prohibited phrases listed in the guidelines (outside quotes).\n"
        "3. Replace each prohibited phrase with its EXACT specified replacement, preserving ال التعريف when required.\n"
        "4. Double-check that titles like 'السلطان' and 'الملك' always include ال التعريف.\n"
        "5. Remove all prayer phrases (حفظه الله, رعاه الله, etc.) completely (outside quotes).\n"
        "6. Rewrite the article according to UNA editorial style.\n"
        "6. CRITICAL REQUIREMENT - OUTPUT FORMAT WITH TITLE (MANDATORY):\n"
        "   YOUR OUTPUT MUST START WITH THE ARTICLE TITLE/HEADLINE, THEN THE ARTICLE BODY.\n"
        "\n"
        "   ⚠️ IMPORTANT: The article title is the FIRST LINE of the article (before the city/date line).\n"
        "   Example: 'جلالةُ السلطان المعظم ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية و6 مذكرات تفاهم'\n"
        "   This is the TITLE/HEADLINE - you MUST process it and return it as the FIRST LINE of your output!\n"
        "\n"
        "   PROCESSING THE TITLE:\n"
        "   - Apply the same honorific removal rules to the title\n"
        "   - Remove exaggerations like 'جلالة', 'المعظم', 'حفظه الله'\n"
        "   - Keep official titles like 'خادم الحرمين الشريفين', 'صاحب السمو الملكي'\n"
        "   - Make the title concise and neutral\n"
        "   Example transformation:\n"
        "   ❌ BEFORE: 'جلالةُ السلطان المعظم ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية و6 مذكرات تفاهم'\n"
        "   ✅ AFTER:  'السلطان ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية و6 مذكرات تفاهم'\n"
        "\n"
        "7. PARAGRAPH DIVISION (MANDATORY) - THIS IS ABSOLUTELY CRITICAL:\n"
        "   🚨🚨🚨 YOU MUST PRESERVE THE PARAGRAPH STRUCTURE! 🚨🚨🚨\n"
        "\n"
        "   ⚠️ CRITICAL RULE: The input article already has paragraphs separated by blank lines (\\n\\n).\n"
        "   You MUST maintain this paragraph structure in your output!\n"
        "\n"
        "   DO NOT merge all paragraphs into one continuous text!\n"
        "   DO NOT rewrite the article as a single long paragraph!\n"
        "\n"
        "   REQUIRED FORMAT:\n"
        "   - Each paragraph from the input should remain a separate paragraph in the output\n"
        "   - Separate EVERY paragraph with exactly TWO newlines (\\n\\n) to create blank lines\n"
        "   - Each paragraph should be 2-4 sentences focusing on ONE main idea\n"
        "   - For SHORT news: 1 paragraph only (title + body + closing)\n"
        "   - For LONG news with multiple topics: 2-4 paragraphs maximum\n"
        "\n"
        "   PARAGRAPH STRUCTURE (preserve from input):\n"
        "   • Paragraph 1: Title/Headline\n"
        "   • [BLANK LINE]\n"
        "   • Paragraph 2: Opening with city/date and main announcement\n"
        "   • [BLANK LINE]\n"
        "   • Paragraph 3: Additional details or context\n"
        "   • [BLANK LINE]\n"
        "   • Paragraph 4: More specific information\n"
        "   • [BLANK LINE]\n"
        "   • Paragraph 5: Supporting details\n"
        "   • [BLANK LINE]\n"
        "   • Paragraph 6: More information\n"
        "   • [BLANK LINE]\n"
        "   • Final line: Closing tag '(انتهى)' or source tag 'العُمانية/' on its OWN separate line\n"
        "\n"
        "   ⚠️ CRITICAL: If the input article ends with 'العُمانية/' or similar source tag,\n"
        "   keep it on a SEPARATE final line after a blank line!\n"
        "\n"
        "   🔴 REAL-WORLD COMPLETE EXAMPLE (EXACTLY HOW OUTPUT SHOULD LOOK):\n"
        "\n"
        "   السلطان ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية و6 مذكرات تفاهم\n"
        "\n"
        "   مدريد في 5 نوفمبر (يونا/العُمانية) - شهد السلطان هيثم بن طارق ورئيس الوزراء الإسباني بيدرو سانشيث اليوم في قصر مونكلوا بمدريد مراسم التوقيع على اتفاقية و6 مذكرات تفاهم بين البلدين الصديقين شملت العديد من المجالات، وذلك في إطار زيارة دولة يقوم بها السلطان إلى مملكة إسبانيا.\n"
        "\n"
        "   تمثلت الاتفاقية في الإعفاء المتبادل من التأشيرات لحاملي جوازات السفر الدبلوماسية والخاصة والخدمة بين سلطنة عُمان ومملكة إسبانيا.\n"
        "\n"
        "   شملت مذكرات التفاهم مجالات الثقافة والرياضة، وترويج الاستثمار، والمجالات الزراعية والحيوانية والسمكية والأمن الغذائي، وإدارة وحماية موارد المياه، والطاقة النظيفة، والنقل والبنية الأساسية.\n"
        "\n"
        "   وقع نيابة عن حكومة سلطنة عُمان كل من وزير الخارجية بدر بن حمد البوسعيدي، ووزير التجارة والصناعة وترويج الاستثمار قيس بن محمد اليوسف، ووزير الطاقة والمعادن سالم بن ناصر العوفي.\n"
        "\n"
        "   وعن حكومة مملكة إسبانيا كل من وزير الشؤون الخارجية والاتحاد الأوروبي والتعاون خوسيه مانويل ألباريس، ووزير الاقتصاد والتجارة والشركات كارلوس كويربو، ووزير الزراعة وصيد الأسماك والغذاء لويس بلاناس.\n"
        "\n"
        "   (انتهى)\n"
        "\n"
        "   ☝️ IMPORTANT NOTES:\n"
        "   - Each paragraph is on its own line, with a BLANK LINE between paragraphs!\n"
        "   - The closing tag '(انتهى)' or 'العُمانية/' is on a SEPARATE line at the end!\n"
        "   - This is NOT one continuous block of text!\n"
        "\n"
        "8. Return ONLY the final revised news article in Arabic with THE TITLE FIRST, then PROPERLY SEPARATED PARAGRAPHS. No analysis or commentary.\n\n"
        "⚠️⚠️⚠️ FINAL REMINDER - READ THIS BEFORE OUTPUTTING ⚠️⚠️⚠️\n"
        "🚨 YOUR OUTPUT MUST HAVE MULTIPLE SEPARATE PARAGRAPHS WITH BLANK LINES BETWEEN THEM! 🚨\n"
        "\n"
        "Required structure:\n"
        "1. Line 1: PROCESSED ARTICLE TITLE (with honorifics removed)\n"
        "2. Line 2: BLANK LINE (\\n\\n)\n"
        "3. Line 3: First paragraph (opening with city/date)\n"
        "4. Line 4: BLANK LINE (\\n\\n)\n"
        "5. Line 5: Second paragraph\n"
        "6. Line 6: BLANK LINE (\\n\\n)\n"
        "7. Line 7: Third paragraph\n"
        "8. ...and so on for ALL paragraphs\n"
        "\n"
        "❌ WRONG (everything in one block):\n"
        "العنوان\\n\\nمدريد في 5 نوفمبر - شهد السلطان... تمثلت الاتفاقية... شملت مذكرات... وقع نيابة... وعن حكومة...\n"
        "\n"
        "✅ CORRECT (separate paragraphs with blank lines):\n"
        "العنوان\\n\\nمدريد في 5 نوفمبر - شهد السلطان...\\n\\nتمثلت الاتفاقية...\\n\\nشملت مذكرات...\\n\\nوقع نيابة...\\n\\nوعن حكومة...\\n\\n(انتهى)\n"
        "\n"
        "⚠️ IMPORTANT: The closing tag '(انتهى)' or 'العُمانية/' must be on a SEPARATE final line!\n"
        "\n"
        "If you do NOT include the title first → YOU FAILED!\n"
        "If you merge paragraphs into one continuous text → YOU FAILED!\n"
        "If you do NOT have blank lines between EVERY paragraph → YOU FAILED!\n"
    )

    return [
        {
            "role": "system",
            "content": (
                "You are an Arabic news processor for UNA (Union of News Agencies).\n\n"
                "🔴🔴🔴 القاعدة الأساسية: لا تُعِد صياغة الخبر! 🔴🔴🔴\n"
                "أنت لست كاتباً، أنت مُعالج نصوص. مهمتك تطبيق تعديلات محددة فقط:\n"
                "1. إضافة الجنسية في العنوان\n"
                "2. تنسيق سطر المدينة والوكالة\n"
                "3. إزالة الألقاب المبالغ فيها\n"
                "4. إضافة (انتهى)\n"
                "5. تصحيح الأخطاء الإملائية\n\n"
                "⛔ ممنوع: إعادة الصياغة، إضافة معلومات، تغيير الكلمات، الحشو\n\n"
                "CONTENT VALIDATION - REJECT INAPPROPRIATE OR RANDOM TEXT:\n"
                "Before processing any text, you MUST validate that it is a legitimate news article:\n"
                "1. REJECT immediately if the text is:\n"
                "   - Random, meaningless, or nonsensical text\n"
                "   - Inappropriate, offensive, or harmful content\n"
                "   - Not a news article (e.g., spam, advertisements, personal messages)\n"
                "   - Contains only symbols, numbers without context, or gibberish\n"
                "   - Clearly not related to news or journalism\n"
                "2. If you reject the text, respond ONLY with: 'ERROR: النص المقدم غير مناسب أو غير صالح للمعالجة. يرجى تقديم خبر صحيح.'\n"
                "3. Only proceed with editing if the text is a legitimate, coherent news article.\n\n"
                "⚠️⚠️⚠️ CRITICAL DISTINCTION: OFFICIAL STATE TITLES vs EXAGGERATION ⚠️⚠️⚠️\n\n"
                "🔴 ABSOLUTE RULE YOU MUST UNDERSTAND:\n"
                "Some phrases are OFFICIAL STATE TITLES (like job titles) - these are NOT exaggeration!\n"
                "Other phrases are EXAGGERATED PRAISE - these must be removed!\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "✅ PRESERVE THESE - OFFICIAL STATE TITLES (NEVER REMOVE!):\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "1. ✅ 'خادم الحرمين الشريفين' - OFFICIAL title of Saudi King\n"
                "   This is NOT exaggeration! It's like saying 'President' or 'Prime Minister'!\n"
                "   ⚠️ NEVER EVER remove this phrase! It's his official state designation!\n\n"
                "2. ✅ 'صاحب السمو الملكي' - OFFICIAL Royal Highness title\n"
                "   This is NOT exaggeration! It's the government-recognized protocol title!\n"
                "   ⚠️ NEVER EVER remove this phrase! It's their official rank in the state!\n\n"
                "3. ✅ 'ولي العهد' - Crown Prince (official position)\n"
                "4. ✅ 'رئيس مجلس الوزراء' - Prime Minister (official position)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "❌ REMOVE THESE - EXAGGERATED PRAISE (NOT OFFICIAL!):\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "1. ❌ 'جلالة الملك المعظم' → ✅ 'الملك' (this IS exaggeration!)\n"
                "2. ❌ 'جلالة الملك' → ✅ 'الملك'\n"
                "3. ❌ 'حضرة صاحب الجلالة الملك' → ✅ 'الملك'\n"
                "4. ❌ 'صاحب الجلالة السلطان' → ✅ 'السلطان'\n"
                "5. ❌ 'فخامة الرئيس' → ✅ 'الرئيس'\n"
                "6. ❌ Prayer phrases: 'حفظه الله', 'أيده الله', 'رعاه الله' → DELETE\n"
                "7. ❌ Exaggeration words: 'المعظم', 'الجليل', 'خالص' → DELETE\n\n"
                "🔑 KEY DIFFERENCE:\n"
                "• 'خادم الحرمين الشريفين' = Like saying 'President Obama' → KEEP!\n"
                "• 'صاحب السمو الملكي' = Like saying 'His Royal Highness' → KEEP!\n"
                "• 'جلالة الملك المعظم' = Like saying 'His Glorious Majesty' → REMOVE!\n"
                "• 'حفظه الله' = Prayer/blessing → REMOVE!\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "📚 MANDATORY EXAMPLES - FOLLOW THESE EXACTLY:\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                "Example 1 - Remove exaggeration, keep simple title:\n"
                "Original: 'بعث جلالة الملك المعظم عبد الله الثاني ابن الحسين حفظه الله برقية تهنئة خالصة'\n"
                "After: 'بعث الملك عبد الله الثاني ابن الحسين برقية تهنئة'\n"
                "Removed: جلالة, المعظم, حفظه الله, خالصة\n"
                "Kept: الملك\n\n"
                "Example 2 - KEEP 'صاحب السمو الملكي' (OFFICIAL TITLE!):\n"
                "Original: 'استقبل صاحب السمو الملكي الأمير سلمان بن حمد آل خليفة ولي العهد رئيس مجلس الوزراء حفظه الله'\n"
                "After: 'استقبل صاحب السمو الملكي الأمير سلمان بن حمد آل خليفة ولي العهد رئيس مجلس الوزراء'\n"
                "Removed: حفظه الله ONLY!\n"
                "Kept: صاحب السمو الملكي (OFFICIAL!), الأمير, ولي العهد, رئيس مجلس الوزراء\n"
                "⚠️ Notice: We did NOT remove 'صاحب السمو الملكي' - it's an OFFICIAL title!\n\n"
                "Example 3 - KEEP 'خادم الحرمين الشريفين' (OFFICIAL TITLE!):\n"
                "Original: 'خادم الحرمين الشريفين الملك سلمان بن عبدالعزيز آل سعود حفظه الله أيده الله'\n"
                "After: 'خادم الحرمين الشريفين الملك سلمان بن عبدالعزيز آل سعود'\n"
                "Removed: حفظه الله, أيده الله ONLY!\n"
                "Kept: خادم الحرمين الشريفين (OFFICIAL SAUDI TITLE!), الملك\n"
                "⚠️ Notice: We did NOT remove 'خادم الحرمين الشريفين' - it's the OFFICIAL title!\n\n"
                "Example 4 - Understanding the difference:\n"
                "❌ WRONG: 'صاحب السمو الملكي الأمير' → 'الأمير' (you removed official title!)\n"
                "✅ RIGHT: 'صاحب السمو الملكي الأمير' → 'صاحب السمو الملكي الأمير' (kept it!)\n"
                "❌ WRONG: 'جلالة الملك' → 'جلالة الملك' (you kept exaggeration!)\n"
                "✅ RIGHT: 'جلالة الملك' → 'الملك' (removed exaggeration!)\n\n"
                "🔴 If you remove 'خادم الحرمين الشريفين' or 'صاحب السمو الملكي', YOU FAILED!\n\n"
                "IMPORTANT RULE: Always preserve 'ال' (definite article) when simplifying titles:\n"
                "✅ CORRECT: 'السلطان' (with ال), 'الملك' (with ال)\n"
                "❌ WRONG: 'سلطان' (without ال), 'ملك' (without ال)\n\n"
                "EDITORIAL STYLE GUIDELINES:\n"
                "1. Use only modern formal Arabic language.\n"
                "2. Avoid emotional and exaggerated expressions.\n"
                "3. Maintain objectivity and balance in all headlines and texts.\n"
                "4. Headlines must be concise and neutral (without exclamation marks or promotional words).\n"
                "5. Do not add personal analyses or conclusions unless from an official explicit source.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "🔴 قاعدة الجنسية في العناوين (إلزامية!) 🔴\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "عند ذكر أي مسؤول أو وزارة في العنوان، يجب إضافة الجنسية:\n"
                "• رئيس الوزراء → رئيس الوزراء العراقي/السعودي/المصري\n"
                "• الملك → الملك السعودي/الأردني/المغربي\n"
                "• الخارجية → الخارجية السعودية/المصرية/الإماراتية\n"
                "• وزارة الصحة → وزارة الصحة السعودية\n"
                "• وزير السياحة → وزير السياحة السعودي\n\n"
                "⚠️ استثناء: 'وزارة الحج والعمرة' لا تحتاج جنسية (وحيدة في العالم)\n\n"
                "كيف تحدد الجنسية:\n"
                "• واع=عراقي | واس=سعودي | وام=إماراتي | وكونا=كويتي | بنا=بحريني\n"
                "• بغداد=عراقي | الرياض=سعودي | القاهرة=مصري\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "🔴🔴🔴 قاعدة الاختصار - الأهم! (CRITICAL!) 🔴🔴🔴\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "⚠️ الخبر = عنوان + فقرة واحدة فقط + (انتهى)\n"
                "⚠️ الحد الأقصى: 3-4 أسطر!\n"
                "⚠️ لا تضف معلومات غير موجودة في الأصل!\n"
                "⚠️ لا تكرر اسم المسؤول!\n"
                "⚠️ لا تضف 'وأوضح' أو 'وأكد أهمية' أو 'جاء ذلك'!\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "🔴🔴🔴 قواعد إلزامية - يجب تطبيقها جميعاً! 🔴🔴🔴\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "⛔ القاعدة 1: العنوان في سطر مستقل (إلزامي)\n"
                "⛔ القاعدة 2: كلمة (انتهى) في سطر مستقل (إلزامي)\n"
                "⛔ القاعدة 3: ذكر اسم الوكالة بعد (يونا/) - مثال: (يونا/واع) أو (يونا/الغد) (إلزامي)\n"
                "⛔ القاعدة 4: العنوان لا يتجاوز 12 كلمة (إلزامي)\n"
                "⛔ القاعدة 5: الفقرة لا تزيد عن 50 كلمة (إلزامي)\n"
                "⛔ القاعدة 6: استبدال 'مراسلنا' بـ 'مراسل [الوكالة المصدر]' - ليس يونا! (إلزامي)\n"
                "⛔ القاعدة 7: التنويع في الأفعال الإعلامية - ممنوع تكرار نفس الفعل (إلزامي)\n"
                "   الأفعال: قال، ذكر، أوضح، أكد، شدّد، بيّن، أشار، أفاد، أعلن، كشف، لفت، صرّح، أبرز\n"
                "⛔ القاعدة 8: تماسك الجمل بأدوات الربط المناسبة (إلزامي)\n"
                "⛔ القاعدة 9: في أخبار فلسطين: استخدم 'استشهاد' بدلاً من 'مقتل/مصرع/موت' (إلزامي)\n"
                "⛔ القاعدة 10: التاريخ بالميلادي فقط: الأحد 21 ديسمبر 2025م (إلزامي)\n"
                "⛔ القاعدة 11: ممنوع الصفات التقريرية بدون دليل: ضخم، تاريخي، غير مسبوق (إلزامي)\n"
                "⛔ القاعدة 12: تجنب الإفراط في الأقواس (إلزامي)\n"
                "⛔ القاعدة 13: ممنوع التكرار غير المبرر (إلزامي)\n"
                "⛔ القاعدة 14: الاهتمام بعلامات الترقيم (الفواصل والنقاط) (إلزامي)\n\n"
                "🚨 تحذير: عدم تطبيق أي قاعدة = فشل! 🚨\n\n"
                "GENERAL EDITING INSTRUCTIONS:\n"
                "1. ⛔ DO NOT rewrite the news - keep the original text!\n"
                "2. Only add nationality to titles in headlines\n"
                "3. Only format the city/agency line\n"
                "4. Only remove exaggerated honorifics\n"
                "5. Only fix spelling errors\n"
                "6. Add (انتهى) at the end\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "📋 STEP-BY-STEP EDITING PROCESS:\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                "Step 1 - Scan for OFFICIAL TITLES to PRESERVE:\n"
                "   ✅ Do you see 'خادم الحرمين الشريفين'? → KEEP IT!\n"
                "   ✅ Do you see 'صاحب السمو الملكي'? → KEEP IT!\n"
                "   ✅ Do you see 'ولي العهد'? → KEEP IT!\n"
                "   ✅ Do you see 'رئيس مجلس الوزراء'? → KEEP IT!\n"
                "   These are OFFICIAL STATE TITLES - like job titles - NOT exaggeration!\n\n"
                "Step 2 - Scan for EXAGGERATION to REMOVE:\n"
                "   ❌ Do you see 'جلالة الملك'? → Change to 'الملك'\n"
                "   ❌ Do you see 'صاحب الجلالة السلطان'? → Change to 'السلطان'\n"
                "   ❌ Do you see 'حفظه الله' or 'أيده الله'? → DELETE completely\n"
                "   ❌ Do you see 'المعظم', 'الجليل', 'خالص'? → DELETE\n"
                "   These are EXAGGERATED PRAISE - not official titles!\n\n"
                "Step 3 - Apply changes CAREFULLY (but preserve quoted text):\n"
                "   🚨 CRITICAL: DO NOT modify any text inside quotation marks!\n"
                "   Text within quotes (\"...\", «...», \"...\") must remain UNCHANGED!\n"
                "   Example: \"جلالة الملك المعظم\" in quotes stays exactly as is!\n"
                "   \n"
                "   For non-quoted text:\n"
                "   • NEVER remove: خادم الحرمين الشريفين, صاحب السمو الملكي\n"
                "   • ALWAYS remove: جلالة, صاحب الجلالة, فخامة\n"
                "   • DELETE prayer phrases: حفظه الله, أيده الله, رعاه الله\n"
                "   • When simplifying, preserve ال: السلطان (not سلطان), الملك (not ملك)\n\n"
                "Step 4 - Final verification checklist:\n"
                "   ✅ Is 'خادم الحرمين الشريفين' still there (if it was in original)?\n"
                "   ✅ Is 'صاحب السمو الملكي' still there (if it was in original)?\n"
                "   ✅ Are all prayer phrases ('حفظه الله', etc.) deleted?\n"
                "   ✅ Are all exaggerations ('جلالة', 'المعظم', etc.) removed?\n"
                "   ✅ Is ال التعريف preserved in simplified titles?\n\n"
                "Step 5 - Apply editorial style guidelines (objectivity, clarity, etc.).\n\n"
                "Step 6 - Rewrite the article according to UNA editorial style.\n\n"
                "8. MANDATORY OUTPUT FORMAT - TITLE FIRST, THEN SEPARATE PARAGRAPHS:\n"
                "   🚨🚨🚨 CRITICAL: PRESERVE PARAGRAPH STRUCTURE WITH BLANK LINES! 🚨🚨🚨\n"
                "\n"
                "   a) IDENTIFY THE TITLE: The title is the FIRST LINE of the input article (before city/date).\n"
                "      Example input title: 'جلالةُ السلطان المعظم ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية'\n"
                "\n"
                "   b) PROCESS THE TITLE: Apply honorific rules to clean the title:\n"
                "      - Remove: 'جلالة', 'المعظم', 'حفظه الله', 'صاحب الجلالة'\n"
                "      - Keep: 'خادم الحرمين الشريفين', 'صاحب السمو الملكي'\n"
                "      Example output title: 'السلطان ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية'\n"
                "\n"
                "   c) OUTPUT FORMAT - MULTIPLE SEPARATE PARAGRAPHS:\n"
                "      Line 1: Processed title\n"
                "      Line 2: BLANK LINE (\\n\\n)\n"
                "      Line 3: First body paragraph (city/date + main news)\n"
                "      Line 4: BLANK LINE (\\n\\n)\n"
                "      Line 5: Second body paragraph\n"
                "      Line 6: BLANK LINE (\\n\\n)\n"
                "      Line 7: Third body paragraph\n"
                "      ...and so on with BLANK LINES between EVERY paragraph\n"
                "\n"
                "   d) PARAGRAPH REQUIREMENTS:\n"
                "      - DO NOT merge paragraphs into one continuous block!\n"
                "      - Each paragraph = 2-4 sentences on ONE topic\n"
                "      - Separate EVERY paragraph with \\n\\n (blank line)\n"
                "      - For SHORT news: 1 paragraph only! For LONG news: 2-4 paragraphs max\n"
                "      - The input article already has paragraph breaks - PRESERVE THEM!\n"
                "      - The closing tag '(انتهى)' or 'العُمانية/' must be on a SEPARATE final line!\n"
                "\n"
                "   e) COMPLETE REAL-WORLD EXAMPLE:\n"
                "      السلطان ورئيس الوزراء الإسباني يشهدان توقيع اتفاقية و6 مذكرات تفاهم\n"
                "      \n"
                "      مدريد في 5 نوفمبر (يونا/العُمانية) - شهد السلطان هيثم بن طارق ورئيس الوزراء...\n"
                "      \n"
                "      تمثلت الاتفاقية في الإعفاء المتبادل من التأشيرات...\n"
                "      \n"
                "      شملت مذكرات التفاهم مجالات الثقافة والرياضة...\n"
                "      \n"
                "      وقع نيابة عن حكومة سلطنة عُمان كل من وزير الخارجية...\n"
                "      \n"
                "      وعن حكومة مملكة إسبانيا كل من وزير الشؤون الخارجية...\n"
                "      \n"
                "      (انتهى)\n"
                "      \n"
                "      ⚠️ NOTE: The closing tag '(انتهى)' or source tag like 'العُمانية/' must be on a SEPARATE final line!\n"
                "\n"
                "9. Deliver ONLY the final revised article in Arabic: TITLE first, then SEPARATE PARAGRAPHS with BLANK LINES between them; no analysis or explanations."
            ),
        },
        {"role": "user", "content": user_prompt},
    ]


def _split_into_paragraphs(text: str) -> str:
    """
    Post-process the AI output to ensure proper paragraph separation.
    This function intelligently splits text into paragraphs based on content patterns.
    """
    # If the text already has proper paragraph breaks, return as is
    if "\n\n" in text and text.count("\n\n") >= 3:
        return text

    # Remove any existing single newlines (but preserve double newlines if they exist)
    text = text.replace("\n\n", "<<<PARAGRAPH_BREAK>>>")
    text = text.replace("\n", " ")
    text = text.replace("<<<PARAGRAPH_BREAK>>>", "\n\n")

    # Step 1: Separate the title from the body
    # Look for city/agency patterns like "بغداد (يونا/واع)" or "الرياض (يونا/واس)"
    city_agency_pattern = r'^(.+?)\s+((?:مدريد|المنامة|الرياض|عمّان|عمان|القاهرة|دمشق|بغداد|مسقط|الكويت|الدوحة|أبوظبي|أبو ظبي|بيروت|تونس|الجزائر|الرباط|طرابلس|نواكشوط|صنعاء|الخرطوم)(?:\s+في\s+\d|\s*\(يونا))'
    match = re.search(city_agency_pattern, text)

    if match:
        # Extract title and body
        title = match.group(1).strip()
        body = text[match.start(2):].strip()
        # Make sure title doesn't end with the city name
        if not any(city in title[-20:] for city in ['بغداد', 'الرياض', 'القاهرة', 'دمشق', 'عمان', 'مسقط']):
            text = f"{title}\n\n{body}"

    # Step 2: Split by common paragraph starters
    paragraph_starters = [
        r'([.؟!])\s+(وأشاد)',
        r'([.؟!])\s+(وأكد)',
        r'([.؟!])\s+(وقال)',
        r'([.؟!])\s+(وأضاف)',
        r'([.؟!])\s+(جاء ذلك)',
        r'([.؟!])\s+(وجاء)',
        r'([.؟!])\s+(كما)',
        r'([.؟!])\s+(وشهد)',
        r'([.؟!])\s+(وشملت)',
        r'([.؟!])\s+(وتمثّلت)',
        r'([.؟!])\s+(تمثلت)',
        r'([.؟!])\s+(شملت)',
        r'([.؟!])\s+(ووقع)',
        r'([.؟!])\s+(وقّع)',
        r'([.؟!])\s+(وقع)',
        r'([.؟!])\s+(وعن)',
        r'([.؟!])\s+(من جانبه)',
        r'([.؟!])\s+(من جهته)',
        r'([.؟!])\s+(من جانبها)',
        r'([.؟!])\s+(بدوره)',
    ]

    # Apply paragraph splitting
    for pattern in paragraph_starters:
        text = re.sub(pattern, r'\1\n\n\2', text)

    # Step 3: Ensure closing tags are on separate lines
    text = re.sub(r'([.؟!])\s*(\(انتهى\))', r'\1\n\n\2', text)
    text = re.sub(r'([.؟!])\s*(العُمانية/)', r'\1\n\n\2', text)

    # Handle case where closing tag is at the end without punctuation
    text = re.sub(r'(\S)\s+(\(انتهى\))', r'\1\n\n\2', text)
    text = re.sub(r'(\S)\s+(العُمانية/)', r'\1\n\n\2', text)

    # Step 4: Clean up any triple or more newlines
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text.strip()


def _search_google_for_fact_check(query: str) -> str:
    """
    Search Google using SERP API to gather information for fact-checking.

    Args:
        query: The search query

    Returns:
        Formatted search results as a string
    """
    serpapi_key = settings.SERPAPI_KEY
    if not serpapi_key:
        # If no SERP API key, return empty results
        return "لا توجد نتائج بحث متاحة."

    try:
        # Remove quotes from the API key if present
        serpapi_key = serpapi_key.strip('"\'')

        params = {
            "q": query,
            "api_key": serpapi_key,
            "num": 5,  # Get top 5 results
            "hl": "ar",  # Arabic language
            "gl": "sa",  # Saudi Arabia region
        }

        search = GoogleSearch(params)
        results = search.get_dict()

        # Extract organic results
        organic_results = results.get("organic_results", [])

        if not organic_results:
            return "لم يتم العثور على نتائج بحث."

        # Format results
        formatted_results = []
        for idx, result in enumerate(organic_results[:5], 1):
            title = result.get("title", "")
            snippet = result.get("snippet", "")
            formatted_results.append(f"{idx}. {title}\n   {snippet}")

        return "\n\n".join(formatted_results)

    except Exception as e:
        # If search fails, return error message
        return f"حدث خطأ أثناء البحث: {str(e)}"


def check_and_correct_text_between_hashtags(
    *,
    text: str,
    model: str = DEFAULT_COMPLETION_MODEL,
    full_context: str = None,
) -> str:
    """
    Extract text between ##text## markers and check/correct factual and linguistic errors using OpenAI.

    Args:
        text: The input text containing ##text## markers
        model: OpenAI model to use for correction
        full_context: Optional full context text to help with fact-checking

    Returns:
        The corrected text that was between the hashtags

    Raises:
        ValueError: If no text is found between ## markers
    """
    # Extract text between ##text##
    pattern = r'##(.+?)##'
    matches = re.findall(pattern, text, re.DOTALL)

    if not matches:
        raise ValueError("No text found between ## markers. Please use format: ##your text here##")

    # Use the first match if multiple exist
    text_to_check = matches[0].strip()

    if not text_to_check:
        raise ValueError("Empty text found between ## markers.")

    # Get context around the marked text for better understanding
    context_text = full_context if full_context else text

    # Search Google for fact-checking with better query
    # Build search query based on context
    search_query = text_to_check

    # Check if this is about a person with a title/position
    if "يونا" in context_text or "UNA" in context_text.upper():
        search_query = "المدير العام اتحاد وكالات أنباء يونا OIC UNA director general 2025"
    elif "المدير" in context_text or "الرئيس" in context_text or "الوزير" in context_text:
        # Extract organization/context from the full text
        search_query = f"{context_text[:200]} من هو"
    else:
        search_query = f"{text_to_check} حقيقة تحقق"

    search_results = _search_google_for_fact_check(search_query)

    client = _get_openai_client()

    # Build the prompt for factual and linguistic correction
    messages = [
        {
            "role": "system",
            "content": (
                "أنت محرر صحفي محترف متخصص في التحقق من صحة المعلومات وتصحيح الأخطاء الواقعية.\n\n"
                "⚠️ مهمتك الرئيسية: التحقق من صحة المعلومات الواقعية وتصحيحها باستخدام نتائج البحث المقدمة\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "1. **استخدام نتائج البحث (الأولوية القصوى):**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                "🔴 سيتم تزويدك بنتائج بحث من جوجل عن المعلومات المطلوب التحقق منها.\n"
                "استخدم هذه النتائج للتحقق من صحة المعلومات وتصحيحها.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "2. **التحقق من المعلومات الجغرافية (مهم جداً):**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                "🔴 قاعدة حاسمة: إذا رأيت كلمة 'العاصمة' أو 'عاصمة'، تحقق فوراً من صحة اسم المدينة!\n\n"
                "عواصم الدول العربية والإسلامية (للمرجعية):\n"
                "• مصر → القاهرة (ليست: الإسكندرية، الجيزة، أسوان)\n"
                "• السعودية → الرياض (ليست: جدة، مكة، المدينة، الطائف)\n"
                "• الإمارات → أبوظبي (ليست: دبي، الشارقة)\n"
                "• الكويت → الكويت\n"
                "• البحرين → المنامة\n"
                "• قطر → الدوحة\n"
                "• عُمان → مسقط\n"
                "• الأردن → عمّان\n"
                "• لبنان → بيروت\n"
                "• سوريا → دمشق\n"
                "• العراق → بغداد\n"
                "• اليمن → صنعاء\n"
                "• فلسطين → القدس (رام الله إدارياً)\n"
                "• السودان → الخرطوم\n"
                "• المغرب → الرباط (ليست: الدار البيضاء)\n"
                "• الجزائر → الجزائر\n"
                "• تونس → تونس\n"
                "• ليبيا → طرابلس\n"
                "• موريتانيا → نواكشوط\n"
                "• الصومال → مقديشو\n"
                "• جيبوتي → جيبوتي\n"
                "• جزر القمر → موروني\n\n"
                "🔴 أمثلة واقعية للتصحيح:\n"
                "❌ 'في العاصمة الطائف' → ✅ 'في الرياض' (الطائف ليست العاصمة!)\n"
                "❌ 'في العاصمة جدة' → ✅ 'في الرياض' (جدة ليست العاصمة!)\n"
                "❌ 'العاصمة المصرية الجيزة' → ✅ 'العاصمة المصرية القاهرة'\n"
                "❌ 'عاصمة الإمارات دبي' → ✅ 'عاصمة الإمارات أبوظبي'\n"
                "❌ 'العاصمة المغربية الدار البيضاء' → ✅ 'العاصمة المغربية الرباط'\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "3. **التحقق من الأسماء والمناصب (مهم جداً!):**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                "🔴 قاعدة حاسمة: إذا رأيت اسم شخص مع منصب (مدير، رئيس، وزير، إلخ)، تحقق فوراً من صحة الاسم!\n\n"
                "استخدم نتائج البحث للتحقق من:\n"
                "   - هل الاسم المذكور هو الشخص الصحيح الذي يشغل هذا المنصب؟\n"
                "   - إذا كان الاسم خاطئاً، استبدله بالاسم الصحيح من نتائج البحث\n\n"
                "أمثلة:\n"
                "   ❌ 'المدير العام ليونا، اسلام بدران' → ✅ 'المدير العام ليونا، محمد بن عبدربه اليامي'\n"
                "   (إذا أظهرت نتائج البحث أن المدير الصحيح هو محمد بن عبدربه اليامي)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "4. **التحقق من معلومات أخرى:**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "   - تحقق من صحة المعلومات التاريخية (تواريخ، أحداث)\n"
                "   - تحقق من صحة الأرقام والإحصائيات إذا كانت واضحة الخطأ\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "5. **التصحيح اللغوي (ثانوي):**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "   - تصحيح الأخطاء الإملائية\n"
                "   - تصحيح الأخطاء النحوية\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "6. **قواعد الإخراج (مهم جداً):**\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "   ✅ أعد فقط النص المصحح بدون أي إضافات\n"
                "   ✅ لا تكتب 'النص المصحح:' أو أي عناوين\n"
                "   ✅ لا تضف شروحات أو تعليقات\n"
                "   ✅ حافظ على أسلوب النص الأصلي\n"
                "   ✅ لا تضف معلومات جديدة\n\n"
                "🔴 تذكير أخير: استخدم نتائج البحث المقدمة للتحقق من صحة المعلومات!"
            ),
        },
        {
            "role": "user",
            "content": (
                f"النص المطلوب التحقق منه وتصحيحه:\n{text_to_check}\n\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"نتائج البحث من جوجل للتحقق من صحة المعلومات:\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                f"{search_results}\n\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"استخدم نتائج البحث أعلاه للتحقق من صحة المعلومات وتصحيحها.\n"
                f"أعد النص المصحح فقط بدون أي شروحات."
            ),
        },
    ]

    result_text = ""
    if hasattr(client, "responses"):
        response = client.responses.create(model=model, input=messages)
        result_text = response.output_text.strip()
    else:
        chat_client = getattr(client, "chat", None)
        if chat_client and hasattr(chat_client, "completions"):
            response = chat_client.completions.create(model=model, messages=messages)
            if response.choices:
                result_text = response.choices[0].message.content.strip()
        else:
            raise RuntimeError("OpenAI client does not support responses or chat completions API.")

    return result_text


def generate_review(
    *,
    news_text: str,
    guideline_chunks: Iterable[RetrievedChunk],
    example_chunks: Iterable[RetrievedChunk],
    model: str = DEFAULT_COMPLETION_MODEL,
) -> str:
    # First, check if there are any ##text## markers and correct them
    pattern = r'##(.+?)##'
    matches = re.findall(pattern, news_text, re.DOTALL)

    if matches:
        # Process each match and replace in the original text
        processed_text = news_text
        for match in matches:
            corrected = check_and_correct_text_between_hashtags(
                text=f"##{match}##",
                model=model,
                full_context=news_text  # Pass full context for better fact-checking
            )
            # Replace the ##original## with the corrected text (without ##)
            processed_text = processed_text.replace(f"##{match}##", corrected, 1)
        news_text = processed_text

    client = _get_openai_client()
    messages = build_review_prompt(news_text, guideline_chunks, example_chunks)

    result_text = ""
    if hasattr(client, "responses"):
        response = client.responses.create(model=model, input=messages)
        result_text = response.output_text.strip()
    else:
        chat_client = getattr(client, "chat", None)
        if chat_client and hasattr(chat_client, "completions"):
            response = chat_client.completions.create(model=model, messages=messages)
            if response.choices:
                result_text = response.choices[0].message.content.strip()
        else:
            raise RuntimeError("OpenAI client does not support responses or chat completions API.")

    # Check if the model rejected the text as inappropriate or random
    if result_text.startswith("ERROR:") or "غير مناسب" in result_text or "غير صالح" in result_text:
        error_msg = result_text.replace("ERROR:", "").strip()
        if not error_msg:
            error_msg = "النص المقدم غير مناسب أو غير صالح للمعالجة. يرجى تقديم خبر صحيح."
        raise ValueError(error_msg)

    # Final pass: ensure any remaining honorifics are processed
    # This catches any honorifics the model might have missed
    final_text = _preprocess_honorifics(result_text)

    # Post-process to ensure proper paragraph separation
    final_text = _split_into_paragraphs(final_text)

    return final_text


# =============================================================================
# DOCUMENT RETRIEVAL FUNCTIONS
# =============================================================================

def list_documents(
    document_type: DocumentChunk.DocumentType | None = None,
) -> dict:
    """
    List all uploaded documents grouped by title.

    Args:
        document_type: Optional filter by document type (guideline/example)

    Returns:
        Dictionary with document summaries
    """
    queryset = DocumentChunk.objects.all()
    if document_type:
        queryset = queryset.filter(document_type=document_type)

    # Group by title and document_type
    from django.db.models import Count, Min

    documents = (
        queryset
        .values('title', 'source_name', 'document_type')
        .annotate(
            total_chunks=Count('id'),
            created_at=Min('created_at')
        )
        .order_by('-created_at')
    )

    return {
        "total_documents": len(documents),
        "total_chunks": queryset.count(),
        "documents": list(documents),
    }


def get_document_detail(
    title: str,
    document_type: DocumentChunk.DocumentType | None = None,
) -> dict | None:
    """
    Get detailed information about a specific document including all chunks.

    Args:
        title: The document title
        document_type: Optional filter by document type

    Returns:
        Dictionary with document details and chunks, or None if not found
    """
    queryset = DocumentChunk.objects.filter(title=title)
    if document_type:
        queryset = queryset.filter(document_type=document_type)

    chunks = list(queryset.order_by('order'))
    if not chunks:
        return None

    first_chunk = chunks[0]
    return {
        "title": first_chunk.title,
        "source_name": first_chunk.source_name,
        "document_type": first_chunk.document_type,
        "total_chunks": len(chunks),
        "created_at": first_chunk.created_at,
        "chunks": [
            {
                "id": chunk.id,
                "document_type": chunk.document_type,
                "title": chunk.title,
                "source_name": chunk.source_name,
                "order": chunk.order,
                "text": chunk.text,
                "metadata": chunk.metadata,
                "created_at": chunk.created_at,
            }
            for chunk in chunks
        ],
    }


def delete_document(
    title: str,
    document_type: DocumentChunk.DocumentType | None = None,
) -> int:
    """
    Delete a document and all its chunks.

    Args:
        title: The document title
        document_type: Optional filter by document type

    Returns:
        Number of chunks deleted
    """
    queryset = DocumentChunk.objects.filter(title=title)
    if document_type:
        queryset = queryset.filter(document_type=document_type)

    count = queryset.count()
    queryset.delete()
    return count


def list_batches(
    document_type: str | None = None,
    status: str | None = None,
) -> List[dict]:
    """
    List all batch uploads with optional filters.

    Args:
        document_type: Optional filter by document type
        status: Optional filter by status

    Returns:
        List of batch records
    """
    queryset = FileUploadBatch.objects.all()
    if document_type:
        queryset = queryset.filter(document_type=document_type)
    if status:
        queryset = queryset.filter(status=status)

    return [
        {
            "batch_id": str(batch.id),
            "document_type": batch.document_type,
            "status": batch.status,
            "total_files": batch.total_files,
            "processed_files": batch.processed_files,
            "total_chunks_created": batch.total_chunks_created,
            "error_message": batch.error_message,
            "created_at": batch.created_at,
            "updated_at": batch.updated_at,
        }
        for batch in queryset
    ]


def get_batch_files(batch_id: str) -> List[dict]:
    """
    Get all files in a batch upload.

    Args:
        batch_id: UUID of the batch

    Returns:
        List of file records
    """
    try:
        batch = FileUploadBatch.objects.get(id=batch_id)
    except FileUploadBatch.DoesNotExist:
        raise ValueError(f"Batch {batch_id} not found")

    return [
        {
            "id": f.id,
            "filename": f.filename,
            "title": f.title,
            "file_size": f.file_size,
            "status": f.status,
            "chunks_created": f.chunks_created,
            "error_message": f.error_message,
            "created_at": f.created_at,
        }
        for f in batch.files.all()
    ]


def get_statistics() -> dict:
    """
    Get overall statistics about uploaded documents.

    Returns:
        Dictionary with statistics
    """
    total_guidelines = DocumentChunk.objects.filter(
        document_type=DocumentChunk.DocumentType.GUIDELINE
    ).count()
    total_examples = DocumentChunk.objects.filter(
        document_type=DocumentChunk.DocumentType.EXAMPLE
    ).count()

    # Count unique documents by title
    guideline_docs = (
        DocumentChunk.objects
        .filter(document_type=DocumentChunk.DocumentType.GUIDELINE)
        .values('title')
        .distinct()
        .count()
    )
    example_docs = (
        DocumentChunk.objects
        .filter(document_type=DocumentChunk.DocumentType.EXAMPLE)
        .values('title')
        .distinct()
        .count()
    )

    return {
        "guidelines": {
            "total_documents": guideline_docs,
            "total_chunks": total_guidelines,
        },
        "examples": {
            "total_documents": example_docs,
            "total_chunks": total_examples,
        },
        "total_documents": guideline_docs + example_docs,
        "total_chunks": total_guidelines + total_examples,
    }

